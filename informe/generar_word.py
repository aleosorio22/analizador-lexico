"""
Genera el informe del Analizador Léxico en formato .docx
Versión actualizada: énfasis en implementación con Expresiones Regulares
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colores ──────────────────────────────────────────────────────────
AZUL       = RGBColor(0x1A, 0x1A, 0x2E)
AZUL_VIV   = RGBColor(0x3B, 0x82, 0xF6)
VIOLETA    = RGBColor(0x8B, 0x5C, 0xF6)
GRIS       = RGBColor(0x44, 0x44, 0x44)
GRIS_SUAVE = RGBColor(0x88, 0x88, 0x88)
VERDE      = RGBColor(0x06, 0x5F, 0x46)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)
ROJO       = RGBColor(0xEF, 0x44, 0x44)
NEGRO      = RGBColor(0x1A, 0x1A, 0x1A)
NARANJA    = RGBColor(0xEA, 0x58, 0x0C)

MONO_FONT = "Courier New"
BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"

# ── Helpers ──────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = HEAD_FONT; run.font.bold = True
    run.font.color.rgb = color or AZUL
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '3B82F6')
        pBdr.append(bot); pPr.append(pBdr)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = BODY_FONT; run.font.size = Pt(11); run.font.color.rgb = NEGRO
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = BODY_FONT; run.font.size = Pt(11); run.font.color.rgb = NEGRO
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_code(text, highlight_color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = MONO_FONT; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC9, 0xD1, 0xD9)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), highlight_color or '0D1117')
    pPr.append(shd)
    return p

def add_code_comment(text):
    """Línea de código resaltada como comentario (color diferente)"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = MONO_FONT; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)   # gris para comentarios
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '0D1117')
    pPr.append(shd)
    return p

def add_callout(text, kind='info'):
    colors = {'info': 'EFF6FF', 'success': 'F0FDF4', 'warning': 'FFF7ED', 'error': 'FEF2F2'}
    txcols = {
        'info':    RGBColor(0x1E, 0x40, 0xAF),
        'success': RGBColor(0x14, 0x53, 0x2D),
        'warning': RGBColor(0x7C, 0x2D, 0x12),
        'error':   RGBColor(0x7F, 0x1D, 0x1D),
    }
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT; run.font.size = Pt(10.5)
    run.font.color.rgb = txcols[kind]
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), colors[kind])
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; set_cell_bg(cell, '1A1A2E')
        p = cell.paragraphs[0]; run = p.add_run(h)
        run.font.name = HEAD_FONT; run.font.bold = True
        run.font.size = Pt(10); run.font.color.rgb = BLANCO
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            if ri % 2 == 1: set_cell_bg(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                run = p.add_run(val[0]); run.font.bold = val[1] if len(val) > 1 else False
                if len(val) > 2:
                    r,g,b = int(val[2][0:2],16), int(val[2][2:4],16), int(val[2][4:6],16)
                    run.font.color.rgb = RGBColor(r,g,b)
            else:
                run = p.add_run(str(val))
            run.font.name = BODY_FONT; run.font.size = Pt(10)
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("Materia: Compiladores")
run.font.name = BODY_FONT; run.font.size = Pt(12); run.font.color.rgb = GRIS

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("Analizador Léxico")
run.font.name = HEAD_FONT; run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = AZUL

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Visualizador Interactivo de Tokens — Implementado con Expresiones Regulares")
run.font.name = BODY_FONT; run.font.size = Pt(13); run.font.color.rgb = GRIS

doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Repositorio: "); run.font.name = BODY_FONT; run.font.size = Pt(11); run.font.bold = True
run2 = p.add_run("https://github.com/aleosorio22/analizador-lexico")
run2.font.name = MONO_FONT; run2.font.size = Pt(10); run2.font.color.rgb = AZUL_VIV

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Demo en vivo: "); run.font.name = BODY_FONT; run.font.size = Pt(11); run.font.bold = True
run2 = p.add_run("https://aleosorio22.github.io/analizador-lexico/")
run2.font.name = MONO_FONT; run2.font.size = Pt(10); run2.font.color.rgb = RGBColor(0x06,0x5F,0x46)

p.paragraph_format.space_before = Pt(40)
for label, value in [
    ("Estudiante:", "Rene Alejandro Osorio"),
    ("GitHub:",     "@aleosorio22"),
    ("Tecnologia:", "JavaScript  React  Vite"),
    ("Fecha:",      "Marzo 2026"),
]:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pp.add_run(f"{label}  ")
    r1.font.name = BODY_FONT; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = NEGRO
    r2 = pp.add_run(value)
    r2.font.name = BODY_FONT; r2.font.size = Pt(11); r2.font.color.rgb = GRIS

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 1 — ANÁLISIS DEL PROBLEMA
# ════════════════════════════════════════════════════════════════════
add_heading("1. Análisis del Problema")

add_heading("Que se pide?", level=3)
add_body(
    "El enunciado solicita construir un Analizador Léxico: un programa capaz de leer código "
    "fuente escrito en cualquier lenguaje de alto nivel e identificar y clasificar cada elemento "
    "del texto en categorías llamadas tokens."
)
add_body(
    "Un analizador léxico es la primera fase de un compilador. Antes de que el compilador pueda "
    "entender la estructura o el significado de un programa, necesita saber qué tipo de palabras "
    "lo componen: si una secuencia de caracteres es un número, un nombre de variable, un operador, "
    "una palabra reservada, etc."
)

add_heading("Flujo del compilador", level=3)
add_code("  Codigo fuente  -->  Analizador Léxico  -->  Lista de Tokens  -->  Analizador Sintáctico")

add_heading("Elementos a reconocer", level=3)
for b in [
    "Identificadores: nombres de variables/funciones, hasta 10 caracteres, inician con letra.",
    "Numeros enteros sin signo: solo valores entre 0 y 100.",
    "Cadenas de caracteres: texto entre comillas simples o dobles.",
    "Operadores aritméticos: +  -  *  /",
    "Operador de asignación: := (estilo Pascal)",
    "Operadores relacionales: >=  <=  >  <  =  <>",
    "Símbolos delimitadores: { } [ ] ( ) , ; ..",
    "Palabras reservadas: if, else, for, print, int + las 120 permutaciones de 'asdfg'.",
]:
    add_bullet(b)

add_heading("Interpretacion del enunciado", level=3)
add_body(
    "Sobre las 'combinaciones posibles con la cadena asdfg': se interpretó que se refiere a todas "
    "las permutaciones de las 5 letras {a, s, d, f, g}. Matematicamente son 5! = 120 combinaciones "
    "únicas, todas tratadas como palabras reservadas del lenguaje."
)
add_callout(
    "IMPORTANTE: El analizador NO se detiene al encontrar un error. Continúa procesando el resto "
    "del código, permitiendo detectar múltiples errores en una sola pasada.",
    kind='info'
)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 2 — DISEÑO DE LA SOLUCIÓN (REGEX)
# ════════════════════════════════════════════════════════════════════
add_heading("2. Diseño de la Solución con Expresiones Regulares")

add_body(
    "El analizador léxico está implementado usando Expresiones Regulares (regex). Esta es la "
    "misma estrategia que utilizan herramientas profesionales de generación de analizadores "
    "como Flex, JFlex y ANTLR, que la profesora menciona en clase."
)

add_heading("Que es una Expresion Regular en este contexto?", level=3)
add_body(
    "Una expresión regular es un patrón que describe qué secuencia de caracteres forma un "
    "token válido. En lugar de describir manualmente carácter por carácter qué hacer, se "
    "declara el patrón de forma compacta y el motor de regex lo resuelve automáticamente."
)

add_table(
    ["Token", "Expresion Regular", "Que describe"],
    [
        ("Identificador",   "^[a-zA-Z][a-zA-Z0-9]*",   "Empieza con letra, seguida de cero o mas letras/digitos"),
        ("Numero entero",   "^\\d+",                    "Uno o más digitos consecutivos"),
        ("Cadena cerrada",  '^"([^"\\n]*)"',             "Todo entre comillas dobles, sin salto de linea"),
        ("Asignacion",      "^:=",                       "Exactamente los dos caracteres : y ="),
        ("Op. relacional",  "^(>=|<=|<>)",              "Cualquiera de los tres operadores de 2 caracteres"),
        ("Simbolo",         "^[{}\\[\\](),;]",          "Cualquier caracter del conjunto de simbolos"),
        ("Blanco",          "^[ \\t\\r\\n]+",           "Uno o mas espacios, tabs o saltos de linea"),
        ("Comentario //",   "^\\/\\/[^\\n]*",           "Dos barras seguidas de cualquier cosa hasta fin de linea"),
    ],
    col_widths=[3.5, 5, 7.5]
)

add_heading("El ancla ^ — por que es critica", level=3)
add_body(
    "Todas las expresiones regulares del proyecto empiezan con ^ (ancla de inicio). "
    "Esto significa que el patron solo coincide al INICIO del texto que queda por procesar. "
    "Sin el ^, la regex podria coincidir en cualquier parte del texto y saltarse caracteres."
)
add_code('  resto = "x := 50"')
add_code('  /^[a-zA-Z]/  -->  coincide "x" al inicio        OK')
add_code('  /[a-zA-Z]/   -->  coincidira con cualquier letra, incluso saltando := y 5')

add_heading("Estrategia: lista de reglas con prioridad", level=3)
add_body(
    "El corazon del analizador es un array llamado RULES que contiene todas las reglas en "
    "orden de prioridad. En cada posicion del texto se prueban las reglas de arriba hacia "
    "abajo y la PRIMERA que coincide gana. Esto resuelve automaticamente la ambiguedad."
)

add_callout(
    "EJEMPLO DE PRIORIDAD: >= debe reconocerse ANTES que >. Si la regla de > estuviera "
    "primero, al encontrar >= produciria el token > y dejaría = sin procesar. Al poner "
    ">= antes en la lista, se reconoce correctamente como un solo token.",
    kind='warning'
)

add_heading("Los 7 tipos de reglas", level=3)
add_table(
    ["Tag (tipo)", "Que hace cuando la regex coincide"],
    [
        ("skip",    "Descarta el texto coincidido (espacios, comentarios). No produce token."),
        ("token",   "Produce un token directamente. El valor es el texto que coincidio."),
        ("string",  "Produce token STRING. El valor es el grupo de captura (contenido sin comillas)."),
        ("strErr",  "La cadena abrio comilla pero no cerro. Produce token de Error."),
        ("number",  "Valida el rango despues del match. Si 0-100 → INTEGER, si no → Error."),
        ("word",    "Valida longitud y busca en reservadas. Puede producir RESERVED, IDENTIFIER o Error."),
        ("unknown", "Fallback: ningun patron coincidio. Produce Error con el caracter no reconocido."),
    ],
    col_widths=[3, 13]
)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 3 — TOKENS IMPLEMENTADOS
# ════════════════════════════════════════════════════════════════════
add_heading("3. Tokens Implementados")

add_table(
    ["Tipo", "Regex usada", "Ejemplos validos", "Error si..."],
    [
        ("Palabra reservada", "^[a-zA-Z][a-zA-Z0-9]* + lookup",  "if, for, asdfg, fgads",  "—"),
        ("Identificador",     "^[a-zA-Z][a-zA-Z0-9]*",           "x, var1, resultado",      "Longitud > 10"),
        ("Numero entero",     "^\\d+  + validar valor",          "0, 50, 100",              "Valor > 100"),
        ("Cadena",            '^"([^"\\n]*)"',                    '"asdfg", "hola"',         "Sin comilla de cierre"),
        ("Op. aritmetico",    "^[+\\-*\\/]",                     "+, -, *, /",              "—"),
        ("Op. asignacion",    "^:=",                             ":=",                      ": sin ="),
        ("Op. relacional",    "^(>=|<=|<>|[><]|=)",             ">, <=, <>, =",            "—"),
        ("Simbolo",           "^[{}\\[\\](),;]  o  ^\\.\\.","{ } [ ] ( ) , ; ..",      ". solo"),
        ("Error",             "^[\\s\\S]  (fallback)",           "@ # $ % &",              "—"),
    ],
    col_widths=[3.5, 4.5, 3.5, 3.5]
)

add_heading("Las 120 palabras reservadas de 'asdfg'", level=3)
add_body(
    "Se generan en tiempo de ejecucion calculando todas las permutaciones del conjunto "
    "{a, s, d, f, g} con un algoritmo recursivo. Total: 5! = 120 palabras."
)
for line in [
    "function* permutations(arr) {",
    "  if (arr.length <= 1) { yield arr; return }",
    "  for (let i = 0; i < arr.length; i++) {",
    "    const rest = [...arr.slice(0,i), ...arr.slice(i+1)]",
    "    for (const perm of permutations(rest)) yield [arr[i], ...perm]",
    "  }",
    "}",
]:
    add_code(line)
add_code_comment("// Resultado: asdfg, asdgf, adsfg, adgfs, ... (120 combinaciones unicas)")
add_code("const ASDFG_PERMS = new Set([...permutations([...'asdfg'])].map(p => p.join('')))")

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 4 — HERRAMIENTAS
# ════════════════════════════════════════════════════════════════════
add_heading("4. Herramientas Utilizadas")

add_table(
    ["Herramienta", "Version", "Rol en el proyecto"],
    [
        ("JavaScript (Regex nativas)", "ES2022+", "Las expresiones regulares son nativas del lenguaje. No se necesita ninguna libreria adicional para el lexer."),
        ("React",          "18.2", "Libreria de interfaz gráfica. Permite construir la UI de forma declarativa con componentes reutilizables."),
        ("Vite",           "5.0",  "Herramienta de build y servidor de desarrollo. Compilacion instantanea y hot-reload."),
        ("CSS puro",       "—",    "Todos los estilos son CSS propio, sin librerias de UI externas."),
        ("GitHub",         "—",    "Repositorio de codigo fuente y control de versiones."),
        ("GitHub Actions", "—",    "CI/CD: construye y publica la app automaticamente en cada push a main."),
        ("GitHub Pages",   "—",    "Hosting gratuito de la aplicacion. URL publica accesible desde cualquier dispositivo."),
    ],
    col_widths=[4, 2, 10]
)

add_callout(
    "NOTA TECNICA: Herramientas como Flex, JFlex y ANTLR que menciona la profesora usan la "
    "misma estrategia: el programador escribe patrones regex, y la herramienta genera "
    "automaticamente el codigo del lexer. En este proyecto se implementa manualmente esa "
    "misma logica, lo que permite entender exactamente como funciona por dentro.",
    kind='info'
)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 5 — ESTRUCTURA DEL PROYECTO
# ════════════════════════════════════════════════════════════════════
add_heading("5. Estructura del Proyecto")

for line in [
    "analizador-lexico/",
    "|-- .github/workflows/deploy.yml   # GitHub Actions: build y deploy automatico",
    "|-- public/favicon.svg             # Icono del compilador (chip con </>)",
    "|-- src/",
    "|   |-- lexer.js                   # * Nucleo: analizador léxico con regex",
    "|   |-- App.jsx                    # Interfaz principal",
    "|   |-- App.css                    # Estilos oscuros tipo VS Code",
    "|   |-- Docs.jsx                   # Modal de documentacion integrada",
    "|   +-- main.jsx                   # Punto de entrada de React",
    "|-- index.html",
    "|-- vite.config.js",
    "|-- package.json",
    "+-- README.md",
]:
    add_code(line)

doc.add_paragraph()
add_table(
    ["Archivo", "Responsabilidad"],
    [
        ("lexer.js",   "Array RULES con todas las regex, clase AnalizadorLexico, 120 permutaciones de 'asdfg', funcion buildSegments() para el resaltado visual."),
        ("App.jsx",    "Componente raiz de la interfaz. Gestiona estado global y renderiza el panel del editor y el panel de tokens."),
        ("Docs.jsx",   "Modal con documentacion completa: tablas de tokens, ejemplos validos/invalidos, programa de ejemplo."),
        ("deploy.yml", "Workflow de GitHub Actions. Se dispara en cada push a main: instala dependencias, compila y publica en GitHub Pages."),
    ],
    col_widths=[3.5, 12.5]
)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 6 — EXPLICACIÓN DEL CÓDIGO (ÉNFASIS REGEX)
# ════════════════════════════════════════════════════════════════════
add_heading("6. Explicacion del Codigo — Enfasis en Expresiones Regulares")

add_body(
    "Esta seccion explica en detalle como funciona cada parte del analizador léxico, "
    "con foco en el uso y la logica de las expresiones regulares."
)

# 6.1 ────────────────────────────────────────────────────────────────
add_heading("6.1 — El array RULES: todas las regex del proyecto", level=3)
add_body(
    "Todas las reglas del lenguaje viven en un solo array llamado RULES. "
    "Cada entrada tiene una regex (re) y una etiqueta (tag) que indica que hacer "
    "cuando la regex coincide. El orden define la prioridad."
)

add_code_comment("// ── 1. Ignorar: blancos y comentarios ──────────────────────────")
add_code("{ tag: 'skip',  re: /^[ \\t\\r\\n]+/       },  // espacios y saltos de linea")
add_code("{ tag: 'skip',  re: /^\\/\\/[^\\n]*/       },  // comentario de linea: //")
add_code("{ tag: 'skip',  re: /^#[^\\n]*/            },  // comentario de linea: #")
add_code("{ tag: 'skip',  re: /^\\/\\*[\\s\\S]*?\\*\\// },  // comentario de bloque: /* */")
add_code(" ")
add_code_comment("// ── 2. Operadores de 2 caracteres (ANTES que los de 1) ──────────")
add_code("{ tag: 'token', re: /^:=/,    type: TK.ASSIGN  },  // asignacion")
add_code("{ tag: 'token', re: /^(>=|<=|<>)/, type: TK.REL_OP },  // relacionales dobles")
add_code("{ tag: 'token', re: /^\\.\\./, type: TK.SYMBOL  },  // rango ..")
add_code(" ")
add_code_comment("// ── 3. Operadores de 1 caracter ─────────────────────────────────")
add_code("{ tag: 'token', re: /^[><]/,  type: TK.REL_OP  },")
add_code("{ tag: 'token', re: /^=/,     type: TK.REL_OP  },")
add_code("{ tag: 'token', re: /^[+\\-*\\/]/, type: TK.ARITH_OP },")
add_code("{ tag: 'token', re: /^[{}\\[\\](),;]/, type: TK.SYMBOL },")
add_code(" ")
add_code_comment("// ── 4. Cadenas (cerrada primero, sin cerrar = error) ─────────────")
add_code('{ tag: "string", re: /^"([^"\\n]*)"/ },  // grupo 1 = contenido sin comillas')
add_code("{ tag: 'string', re: /^'([^'\\n]*)'/ },")
add_code('{ tag: "strErr", re: /^"[^"\\n]*/    },  // abrio pero no cerro')
add_code("{ tag: 'strErr', re: /^'[^'\\n]*/    },")
add_code(" ")
add_code_comment("// ── 5 y 6. Numeros y palabras (logica extra despues del match) ───")
add_code("{ tag: 'number', re: /^\\d+/ },")
add_code("{ tag: 'word',   re: /^[a-zA-Z][a-zA-Z0-9]*/ },")
add_code(" ")
add_code_comment("// ── 7. Fallback: caracter no reconocido ──────────────────────────")
add_code("{ tag: 'unknown', re: /^[\\s\\S]/ },")

doc.add_paragraph()

# 6.2 ────────────────────────────────────────────────────────────────
add_heading("6.2 — El loop principal: como se aplican las regex", level=3)
add_body(
    "El metodo analyze() tiene un loop que se repite mientras quede texto por procesar. "
    "En cada iteracion toma el texto restante y lo pone a prueba contra cada regla:"
)

add_code("analyze() {")
add_code("  while (this.pos < this.src.length) {")
add_code(" ")
add_code_comment("    // Toma el texto que falta procesar desde la posicion actual")
add_code("    const rest  = this.src.slice(this.pos)")
add_code("    const start = this.pos")
add_code(" ")
add_code_comment("    // Prueba cada regla en orden de prioridad")
add_code("    for (const rule of RULES) {")
add_code("      const m = rest.match(rule.re)   // intenta la regex")
add_code("      if (!m) continue                // no coincidio, probar la siguiente")
add_code(" ")
add_code_comment("      // Primera que coincide: procesar segun su tag y salir del for")
add_code("      switch (rule.tag) { ... }")
add_code("      break  // regla aplicada, volver al while con la nueva posicion")
add_code("    }")
add_code("  }")
add_code("}")

doc.add_paragraph()

# 6.3 ────────────────────────────────────────────────────────────────
add_heading("6.3 — Como se procesa cada tipo de regla (switch)", level=3)
add_body(
    "Cuando una regex coincide, el tag determina que hacer con el texto encontrado. "
    "Los casos mas importantes son:"
)

add_heading("TAG: 'token' — el caso mas simple", level=3)
add_body(
    "La regex coincide y el valor del token es exactamente el texto que coincidio. "
    "No se necesita logica extra."
)
add_code_comment("// Ejemplo: rest = ':= 50', regex /^:=/ coincide con ':='")
add_code("case 'token':")
add_code("  this._consume(raw.length)          // avanza 2 posiciones")
add_code("  this._push(rule.type, raw, start)  // agrega token ':=' tipo ASSIGN")
add_code("  break")

add_heading("TAG: 'string' — grupos de captura", level=3)
add_body(
    "Las cadenas usan un GRUPO DE CAPTURA en la regex: los parentesis ( ) capturan "
    "solo el contenido, sin las comillas. El valor del token es m[1], no m[0]."
)
add_code_comment('// Regex: /^"([^"\\n]*)"/')
add_code_comment('// Para el texto: "hola mundo"')
add_code_comment('//   m[0] = \'"hola mundo"\'  (con comillas)')
add_code_comment('//   m[1] = \'hola mundo\'     (sin comillas)  <-- esto es el valor del token')
add_code("case 'string':")
add_code("  this._consume(raw.length)")
add_code("  this._push(TK.STRING, m[1], start)  // m[1] = contenido sin comillas")
add_code("  break")

add_heading("TAG: 'number' — regex + validacion de rango", level=3)
add_body(
    "La regex /^\\d+/ detecta que hay un numero, pero no puede validar que su VALOR "
    "este entre 0 y 100. Eso lo hace el codigo despues del match. Esta es la limitacion "
    "de las regex puras: detectan patrones de texto, no valores semanticos."
)
add_code_comment("// La regex detecto que es un numero: raw = '999'")
add_code("case 'number': {")
add_code("  const val = parseInt(raw, 10)    // convierte el texto a numero")
add_code("  this._consume(raw.length)")
add_code("  if (val >= 0 && val <= 100)      // validacion que regex no puede hacer")
add_code("    this._push(TK.INTEGER, raw, start)")
add_code("  else")
add_code("    this._push(TK.ERROR, `numero fuera de rango: ${raw}`, start)")
add_code("  break")
add_code("}")

add_heading("TAG: 'word' — regex + multiples validaciones", level=3)
add_body(
    "Similar al caso del numero: la regex detecta la palabra, pero hay tres preguntas "
    "que el codigo debe responder porque la regex sola no puede:"
)
add_bullet("¿Tiene mas de 10 caracteres?  →  Error")
add_bullet("¿Esta en la lista de reservadas base (if, else, for...)?  →  Palabra reservada")
add_bullet("¿Es una permutacion de 'asdfg'?  →  Palabra reservada")
add_bullet("Ninguna de las anteriores  →  Identificador")
add_code(" ")
add_code_comment("// La regex detecto una palabra: raw = 'resultado'  (10 chars, exacto)")
add_code("case 'word': {")
add_code("  this._consume(raw.length)")
add_code("  if (raw.length > 10)                        // 1: demasiado largo")
add_code('    this._push(TK.ERROR, `demasiado largo: "${raw}"`, start)')
add_code("  else if (BASE_RESERVED.has(raw.toLowerCase()))  // 2: reservada base")
add_code("    this._push(TK.RESERVED, raw.toLowerCase(), start)")
add_code("  else if (ASDFG_PERMS.has(raw))              // 3: permutacion asdfg")
add_code("    this._push(TK.RESERVED, raw, start)")
add_code("  else                                         // 4: identificador normal")
add_code("    this._push(TK.IDENTIFIER, raw, start)")
add_code("  break")
add_code("}")

doc.add_paragraph()

# 6.4 ────────────────────────────────────────────────────────────────
add_heading("6.4 — _consume(): avanzar contando lineas", level=3)
add_body(
    "Cuando una regex coincide y se consume el texto, hay que actualizar la posicion "
    "actual y contar cuantos saltos de linea habia en el texto consumido. "
    "Esto permite reportar el numero de linea correcto en cada token."
)
add_code("_consume(len) {")
add_code("  const chunk = this.src.slice(this.pos, this.pos + len)")
add_code_comment("  // Contar cuantos '\\n' habia en el texto que acabamos de consumir")
add_code("  this.line += (chunk.match(/\\n/g) || []).length")
add_code("  this.pos  += len")
add_code("}")

doc.add_paragraph()

# 6.5 ────────────────────────────────────────────────────────────────
add_heading("6.5 — Comparacion: que hace la regex vs que hace el codigo", level=3)
add_body(
    "Esta tabla resume la division de responsabilidades entre las expresiones regulares "
    "y la logica adicional en el codigo:"
)
add_table(
    ["Tarea", "Lo hace la regex?", "Como se resuelve"],
    [
        ("Detectar que algo es un numero",          "SI",  "Regex /^\\d+/"),
        ("Validar que el numero este entre 0-100",   "NO",  "parseInt() + if en el codigo"),
        ("Detectar que algo es una palabra",         "SI",  "Regex /^[a-zA-Z][a-zA-Z0-9]*/"),
        ("Validar que tenga max 10 caracteres",      "NO",  "raw.length > 10 en el codigo"),
        ("Saber si es palabra reservada",            "NO",  "BASE_RESERVED.has() o ASDFG_PERMS.has()"),
        ("Detectar cadena correctamente cerrada",    "SI",  "Regex con grupo de captura /^\"([^\"\\n]*)\"/"),
        ("Detectar cadena sin cerrar",               "SI",  "Regex sin comilla final /^\"[^\"\\n]*/"),
        ("Distinguir >= de >",                       "SI",  "Orden de prioridad en RULES (>= antes que >)"),
        ("Ignorar comentarios // y /* */",           "SI",  "Regex /^\\/\\/[^\\n]*/ y /^\\/\\*[\\s\\S]*?\\*\\//"),
        ("Contar numero de linea",                   "NO",  "_consume() cuenta \\n en texto consumido"),
    ],
    col_widths=[7, 2.5, 6.5]
)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 7 — INTERFAZ GRÁFICA
# ════════════════════════════════════════════════════════════════════
add_heading("7. Interfaz Grafica")

add_body(
    "La interfaz esta dividida en tres areas principales con un tema oscuro inspirado "
    "en editores de codigo profesionales como VS Code:"
)
add_bullet("Header: Logo del compilador (chip SVG), titulo, nombre del autor, link a GitHub")
add_bullet("Panel izquierdo: Editor de codigo con numeros de linea y vista de tokens resaltados")
add_bullet("Panel derecho: Lista de tokens con badges de color, filtros por tipo y estadisticas")

add_heading("7.1 — Vista principal", level=3)
add_body("[Insertar captura: 01-interfaz-principal.png — App con el codigo de ejemplo cargado]")

add_heading("7.2 — Resultado del analisis con tokens resaltados", level=3)
add_body("[Insertar captura: 02-tokens-resaltados.png — Codigo con colores y lista de tokens]")

add_heading("Colores por tipo de token", level=3)
add_table(
    ["Color", "Tipo de token"],
    [
        ("Violeta",  "Palabras reservadas"),
        ("Azul",     "Identificadores"),
        ("Verde",    "Numeros enteros"),
        ("Ambar",    "Cadenas de caracteres"),
        ("Naranja",  "Operadores aritmeticos"),
        ("Indigo",   "Operador de asignacion"),
        ("Cian",     "Operadores relacionales"),
        ("Gris",     "Simbolos"),
        ("Rojo",     "Errores lexicos"),
    ],
    col_widths=[4, 12]
)

add_heading("7.3 — Deteccion de errores y filtros", level=3)
add_body("[Insertar captura: 03-errores-filtros.png — Filtro de Error activo]")

add_heading("7.4 — Modal de documentacion", level=3)
add_body("[Insertar captura: 04-documentacion.png — Modal de documentacion abierto]")

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 8 — REPOSITORIO Y DEMO
# ════════════════════════════════════════════════════════════════════
add_heading("8. Repositorio y Demo en Vivo")

add_heading("Repositorio en GitHub", level=3)
p = doc.add_paragraph()
run = p.add_run("https://github.com/aleosorio22/analizador-lexico")
run.font.name = MONO_FONT; run.font.size = Pt(11); run.font.color.rgb = AZUL_VIV; run.font.bold = True

add_heading("Demo en vivo — GitHub Pages", level=3)
p = doc.add_paragraph()
run = p.add_run("https://aleosorio22.github.io/analizador-lexico/")
run.font.name = MONO_FONT; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x06,0x5F,0x46); run.font.bold = True

add_body(
    "La aplicacion esta publicada en GitHub Pages y es accesible desde cualquier navegador "
    "sin necesidad de instalar nada. El deploy es automatico via GitHub Actions: cada push "
    "a la rama main construye y publica la nueva version automaticamente."
)

add_heading("Instrucciones para correr localmente", level=3)
for line in [
    "# Clonar el repositorio",
    "git clone https://github.com/aleosorio22/analizador-lexico.git",
    "cd analizador-lexico",
    " ",
    "# Instalar dependencias",
    "npm install",
    " ",
    "# Iniciar servidor de desarrollo",
    "npm run dev",
    "# Abrir en: http://localhost:5173",
]:
    add_code(line)

page_break()


# ════════════════════════════════════════════════════════════════════
#  SECCIÓN 9 — CONCLUSIONES
# ════════════════════════════════════════════════════════════════════
add_heading("9. Conclusiones")

add_body(
    "El analizador léxico implementado cumple con todos los requerimientos del enunciado "
    "utilizando expresiones regulares como estrategia principal de reconocimiento de tokens. "
    "Esta es la misma tecnica que usan compiladores y herramientas profesionales de la industria."
)

add_body(
    "La implementacion con regex tiene ventajas claras sobre el enfoque caracter a caracter: "
    "el codigo es mas compacto, cada patron es declarativo (describe QUE buscar, no COMO buscarlo), "
    "y la logica de operadores de dos caracteres como >= o := queda resuelta automaticamente "
    "por el orden de prioridad de las reglas."
)

add_heading("Aspectos clave del aprendizaje", level=3)
add_bullet(
    "Las regex detectan PATRONES de texto, pero la logica semantica (rango 0-100, "
    "longitud maxima, lookup en reservadas) siempre requiere codigo adicional."
)
add_bullet(
    "El ancla ^ en cada regex es fundamental: garantiza que el patron solo coincida "
    "al inicio del texto restante, no en cualquier posicion."
)
add_bullet(
    "El ORDEN de las reglas define la prioridad. Poner >= antes que > resuelve "
    "automaticamente la ambiguedad entre operadores de distinto largo."
)
add_bullet(
    "Los grupos de captura ( ) en las regex de cadenas permiten extraer el contenido "
    "sin las comillas de forma elegante, sin necesidad de procesamiento posterior."
)
add_bullet(
    "Herramientas como Flex y ANTLR implementan exactamente esta misma estrategia "
    "pero la generan automaticamente a partir de los patrones que escribe el programador."
)
add_bullet(
    "Las 120 palabras reservadas generadas por permutacion de 'asdfg' demuestran que "
    "el vocabulario de un lenguaje puede definirse matematicamente."
)

add_callout(
    "El proyecto esta disponible en aleosorio22.github.io/analizador-lexico y puede "
    "inspeccionarse, probarse y extenderse para la siguiente etapa de la materia.",
    kind='success'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Rene Alejandro Osorio  |  @aleosorio22  |  Compiladores  |  Marzo 2026")
run.font.name = BODY_FONT; run.font.size = Pt(9.5); run.font.color.rgb = GRIS_SUAVE


# ── Guardar ──────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), "Informe_Analizador_Lexico.docx")
doc.save(out)
print(f"Documento generado: {out}")
